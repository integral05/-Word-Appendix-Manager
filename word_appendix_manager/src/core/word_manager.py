"""
Word Document Manager
Handles all interactions with Microsoft Word documents using COM automation (Windows)
and python-docx (macOS/fallback).
"""

import os
import sys
import platform
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple
from datetime import datetime

# Add parent directory to path for imports
sys.path.append(str(Path(__file__).parent.parent))

from utils.logger import get_logger, LoggerMixin
from utils.exceptions import (
    WordDocumentError, WordNotAvailableError, WordDocumentNotFoundError,
    FileSystemError
)

# Platform-specific imports
if platform.system() == "Windows":
    try:
        import win32com.client as win32
        import pythoncom
        COM_AVAILABLE = True
    except ImportError:
        COM_AVAILABLE = False
else:
    COM_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordManager(LoggerMixin):
    """Manages Microsoft Word document operations."""
    
    def __init__(self, settings=None):
        self.settings = settings
        self.word_app = None
        self.active_document = None
        self.document_path = None
        self.backup_created = False
        
        # Check available backends
        self.com_available = COM_AVAILABLE and platform.system() == "Windows"
        self.docx_available = DOCX_AVAILABLE
        
        if not self.com_available and not self.docx_available:
            raise WordNotAvailableError(
                "Neither COM automation nor python-docx is available. "
                "Please install required dependencies."
            )
        
        self.logger.info(f"WordManager initialized - COM: {self.com_available}, DOCX: {self.docx_available}")
    
    def get_open_documents(self) -> List[Dict[str, str]]:
        """Get list of currently open Word documents (COM only)."""
        if not self.com_available:
            self.logger.warning("COM not available - cannot get open documents")
            return []
        
        try:
            self._ensure_word_app()
            documents = []
            
            for i in range(1, self.word_app.Documents.Count + 1):
                doc = self.word_app.Documents(i)
                doc_info = {
                    "name": doc.Name,
                    "full_name": doc.FullName if doc.Saved else f"{doc.Name} (unsaved)",
                    "path": doc.FullName if doc.Saved else "",
                    "saved": doc.Saved,
                    "index": i
                }
                documents.append(doc_info)
            
            self.logger.info(f"Found {len(documents)} open documents")
            return documents
            
        except Exception as e:
            self.logger.error(f"Failed to get open documents: {e}")
            return []
    
    def open_document(self, document_path: str) -> bool:
        """Open a Word document for editing."""
        document_path = Path(document_path).resolve()
        
        if not document_path.exists():
            raise WordDocumentNotFoundError(str(document_path))
        
        try:
            if self.com_available:
                return self._open_document_com(document_path)
            elif self.docx_available:
                return self._open_document_docx(document_path)
            else:
                raise WordNotAvailableError()
                
        except Exception as e:
            self.logger.error(f"Failed to open document {document_path}: {e}")
            raise WordDocumentError(f"Failed to open document: {str(e)}", str(document_path), "open_document")
    
    def _open_document_com(self, document_path: Path) -> bool:
        """Open document using COM automation."""
        try:
            self._ensure_word_app()
            
            # Check if document is already open
            for doc in self.word_app.Documents:
                if doc.FullName.lower() == str(document_path).lower():
                    self.active_document = doc
                    self.document_path = document_path
                    self.logger.info(f"Document already open: {document_path}")
                    return True
            
            # Open the document
            self.active_document = self.word_app.Documents.Open(str(document_path))
            self.document_path = document_path
            
            self.logger.info(f"Opened document via COM: {document_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"COM document open failed: {e}")
            raise
    
    def _open_document_docx(self, document_path: Path) -> bool:
        """Open document using python-docx."""
        try:
            self.active_document = Document(str(document_path))
            self.document_path = document_path
            
            self.logger.info(f"Opened document via python-docx: {document_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"python-docx document open failed: {e}")
            raise
    
    def select_open_document(self, document_index: int) -> bool:
        """Select an open document by index (COM only)."""
        if not self.com_available:
            raise WordNotAvailableError("Document selection requires COM automation")
        
        try:
            self._ensure_word_app()
            self.active_document = self.word_app.Documents(document_index)
            self.document_path = Path(self.active_document.FullName) if self.active_document.Saved else None
            
            self.logger.info(f"Selected document {document_index}: {self.active_document.Name}")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to select document {document_index}: {e}")
            raise WordDocumentError(f"Failed to select document: {str(e)}", operation="select_document")
    
    def create_backup(self, backup_dir: Optional[str] = None) -> str:
        """Create a backup of the current document."""
        if not self.document_path:
            raise WordDocumentError("No document loaded to backup", operation="create_backup")
        
        try:
            # Determine backup directory
            if backup_dir is None and self.settings:
                backup_dir = self.settings.backup_directory
            if backup_dir is None:
                backup_dir = Path.home() / "Documents" / "WordAppendixManager_Backups"
            
            backup_dir = Path(backup_dir)
            backup_dir.mkdir(parents=True, exist_ok=True)
            
            # Create backup filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            original_name = self.document_path.stem
            backup_name = f"{original_name}_backup_{timestamp}.docx"
            backup_path = backup_dir / backup_name
            
            # Save backup
            if self.com_available and hasattr(self.active_document, 'SaveAs2'):
                self.active_document.SaveAs2(str(backup_path))
                # Reopen original document
                self.active_document = self.word_app.Documents.Open(str(self.document_path))
            else:
                # Copy file manually for python-docx
                import shutil
                shutil.copy2(self.document_path, backup_path)
            
            self.backup_created = True
            self.logger.info(f"Created backup: {backup_path}")
            return str(backup_path)
            
        except Exception as e:
            self.logger.error(f"Failed to create backup: {e}")
            raise FileSystemError(f"Failed to create backup: {str(e)}", str(self.document_path))
    
    def add_appendix_section(self, appendix_title: str, page_count: int, heading_style: Optional[Dict] = None) -> bool:
        """Add an appendix section with blank pages to the document."""
        if not self.active_document:
            raise WordDocumentError("No active document", operation="add_appendix")
        
        try:
            if self.com_available:
                return self._add_appendix_com(appendix_title, page_count, heading_style)
            else:
                return self._add_appendix_docx(appendix_title, page_count, heading_style)
                
        except Exception as e:
            self.logger.error(f"Failed to add appendix {appendix_title}: {e}")
            raise WordDocumentError(f"Failed to add appendix: {str(e)}", operation="add_appendix")
    
    def _add_appendix_com(self, appendix_title: str, page_count: int, heading_style: Optional[Dict] = None) -> bool:
        """Add appendix using COM automation."""
        try:
            # Move to end of document
            selection = self.word_app.Selection
            selection.EndKey(Unit=6)  # wdStory
            
            # Insert page break
            selection.InsertBreak(Type=7)  # wdPageBreak
            
            # Add heading
            selection.TypeText(appendix_title)
            selection.TypeParagraph()
            
            # Apply heading style
            if heading_style:
                self._apply_heading_style_com(selection, heading_style)
            
            # Add blank pages
            for i in range(page_count - 1):
                selection.InsertBreak(Type=7)  # wdPageBreak
                selection.TypeParagraph()
            
            self.logger.info(f"Added appendix '{appendix_title}' with {page_count} pages via COM")
            return True
            
        except Exception as e:
            self.logger.error(f"COM appendix addition failed: {e}")
            raise
    
    def _add_appendix_docx(self, appendix_title: str, page_count: int, heading_style: Optional[Dict] = None) -> bool:
        """Add appendix using python-docx."""
        try:
            # Add page break
            self.active_document.add_page_break()
            
            # Add heading
            heading = self.active_document.add_heading(appendix_title, level=1)
            
            # Apply heading style
            if heading_style:
                self._apply_heading_style_docx(heading, heading_style)
            
            # Add blank pages
            for i in range(page_count - 1):
                self.active_document.add_page_break()
                self.active_document.add_paragraph()
            
            self.logger.info(f"Added appendix '{appendix_title}' with {page_count} pages via python-docx")
            return True
            
        except Exception as e:
            self.logger.error(f"python-docx appendix addition failed: {e}")
            raise
    
    def _apply_heading_style_com(self, selection, heading_style: Dict):
        """Apply heading style using COM."""
        font = selection.Font
        if 'font_name' in heading_style:
            font.Name = heading_style['font_name']
        if 'font_size' in heading_style:
            font.Size = heading_style['font_size']
        if 'bold' in heading_style:
            font.Bold = heading_style['bold']
        if 'italic' in heading_style:
            font.Italic = heading_style['italic']
    
    def _apply_heading_style_docx(self, heading, heading_style: Dict):
        """Apply heading style using python-docx."""
        run = heading.runs[0] if heading.runs else heading.add_run()
        if 'font_name' in heading_style:
            run.font.name = heading_style['font_name']
        if 'font_size' in heading_style:
            run.font.size = Pt(heading_style['font_size'])
        if 'bold' in heading_style:
            run.font.bold = heading_style['bold']
        if 'italic' in heading_style:
            run.font.italic = heading_style['italic']
    
    def save_document(self, file_path: Optional[str] = None) -> bool:
        """Save the document."""
        if not self.active_document:
            raise WordDocumentError("No active document to save", operation="save_document")
        
        try:
            if self.com_available:
                if file_path:
                    self.active_document.SaveAs2(file_path)
                else:
                    self.active_document.Save()
            else:
                save_path = file_path or str(self.document_path)
                self.active_document.save(save_path)
            
            self.logger.info(f"Document saved to: {file_path or 'current location'}")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to save document: {e}")
            raise WordDocumentError(f"Failed to save document: {str(e)}", operation="save_document")
    
    def export_as_pdf(self, pdf_path: str) -> bool:
        """Export document as PDF."""
        if not self.active_document:
            raise WordDocumentError("No active document to export", operation="export_pdf")
        
        try:
            if self.com_available:
                self.active_document.ExportAsFixedFormat(
                    OutputFileName=pdf_path,
                    ExportFormat=17  # wdExportFormatPDF
                )
            else:
                raise WordDocumentError(
                    "PDF export not available with python-docx backend. "
                    "Use COM automation on Windows for full functionality.",
                    operation="export_pdf"
                )
            
            self.logger.info(f"Document exported as PDF: {pdf_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to export PDF: {e}")
            raise WordDocumentError(f"Failed to export PDF: {str(e)}", operation="export_pdf")
    
    def get_document_info(self) -> Dict[str, Any]:
        """Get information about the current document."""
        if not self.active_document:
            return {}
        
        try:
            if self.com_available:
                info = {
                    "name": self.active_document.Name,
                    "path": self.active_document.FullName if self.active_document.Saved else "",
                    "saved": self.active_document.Saved,
                    "page_count": self.active_document.ComputeStatistics(2),
                    "word_count": self.active_document.ComputeStatistics(0),
                    "modified": not self.active_document.Saved
                }
            else:
                info = {
                    "name": self.document_path.name if self.document_path else "Unknown",
                    "path": str(self.document_path) if self.document_path else "",
                    "saved": True,
                    "page_count": len(self.active_document.sections),
                    "word_count": sum(len(p.text.split()) for p in self.active_document.paragraphs),
                    "modified": False
                }
            
            return info
            
        except Exception as e:
            self.logger.error(f"Failed to get document info: {e}")
            return {"error": str(e)}
    
    def _ensure_word_app(self):
        """Ensure Word application is available and connected."""
        if not self.com_available:
            return
        
        try:
            if not self.word_app:
                # Try to connect to existing Word instance
                try:
                    self.word_app = win32.GetActiveObject("Word.Application")
                    self.logger.info("Connected to existing Word instance")
                except:
                    # Create new Word instance
                    self.word_app = win32.Dispatch("Word.Application")
                    self.word_app.Visible = True
                    self.logger.info("Created new Word instance")
            
            # Test connection
            _ = self.word_app.Documents.Count
            
        except Exception as e:
            self.logger.error(f"Failed to ensure Word app: {e}")
            self.word_app = None
            raise WordNotAvailableError(f"Cannot connect to Word application: {str(e)}")
    
    def close(self):
        """Clean up resources."""
        try:
            if self.com_available and self.word_app:
                self.word_app = None
            
            self.active_document = None
            self.document_path = None
            self.logger.info("WordManager closed")
            
        except Exception as e:
            self.logger.error(f"Error during WordManager cleanup: {e}")
    
    def __enter__(self):
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close() 
